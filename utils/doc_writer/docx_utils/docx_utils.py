import markdown
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
import json 
import re 

from docx.text.paragraph import Paragraph
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import markdown2 
from docx import Document
from docx.enum.style import WD_STYLE_TYPE


from bs4 import BeautifulSoup, NavigableString, Tag
from anytree import Node, AsciiStyle, RenderTree, PreOrderIter


from pathlib import Path


STYLES = ['Normal',
 'Header',
 'Header Char',
 'Footer',
 'Footer Char',
 'Heading 1',
 'Heading 2',
 'Heading 3',
 'Heading 4',
 'Heading 5',
 'Heading 6',
 'Heading 7',
 'Heading 8',
 'Heading 9',
 'Default Paragraph Font',
 'Normal Table',
 'No List',
 'No Spacing',
 'Heading 1 Char',
 'Heading 2 Char',
 'Heading 3 Char',
 'Title',
 'Title Char',
 'Subtitle',
 'Subtitle Char',
 'List Paragraph',
 'Body Text',
 'Body Text Char',
 'Body Text 2',
 'Body Text 2 Char',
 'Body Text 3',
 'Body Text 3 Char',
 'List',
 'List 2',
 'List 3',
 'List Bullet',
 'List Bullet 2',
 'List Bullet 3',
  'List Bullet 4',
 'List Bullet 5',
 'List Bullet 6',

 'List Number',
 'List Number 2',
 'List Number 3',
  'List Number 4',
 'List Number 5',
 'List Number 6',

 'List Continue',
 'List Continue 2',
 'List Continue 3',
 'macro',
 'Macro Text Char',
 'Quote',
 'Quote Char',
 'Heading 4 Char',
 'Heading 5 Char',
 'Heading 6 Char',
 'Heading 7 Char',
 'Heading 8 Char',
 'Heading 9 Char',
 'Caption',
 'Strong',
 'Emphasis',
 'Intense Quote',
 'Intense Quote Char',
 'Subtle Emphasis',
 'Intense Emphasis',
 'Subtle Reference',
 'Intense Reference',
 'Book Title',
 'TOC Heading',
 'Table Grid',
 'Light Shading',
 'Light Shading Accent 1',
 'Light Shading Accent 2',
 'Light Shading Accent 3',
 'Light Shading Accent 4',
 'Light Shading Accent 5',
 'Light Shading Accent 6',
 'Light List',
 'Light List Accent 1',
 'Light List Accent 2',
 'Light List Accent 3',
 'Light List Accent 4',
 'Light List Accent 5',
 'Light List Accent 6',
 'Light Grid',
 'Light Grid Accent 1',
 'Light Grid Accent 2',
 'Light Grid Accent 3',
 'Light Grid Accent 4',
 'Light Grid Accent 5',
 'Light Grid Accent 6',
 'Medium Shading 1',
 'Medium Shading 1 Accent 1',
 'Medium Shading 1 Accent 2',
 'Medium Shading 1 Accent 3',
 'Medium Shading 1 Accent 4',
 'Medium Shading 1 Accent 5',
 'Medium Shading 1 Accent 6',
 'Medium Shading 2',
 'Medium Shading 2 Accent 1',
 'Medium Shading 2 Accent 2',
 'Medium Shading 2 Accent 3',
 'Medium Shading 2 Accent 4',
 'Medium Shading 2 Accent 5',
 'Medium Shading 2 Accent 6',
 'Medium List 1',
 'Medium List 1 Accent 1',
 'Medium List 1 Accent 2',
 'Medium List 1 Accent 3',
 'Medium List 1 Accent 4',
 'Medium List 1 Accent 5',
 'Medium List 1 Accent 6',
 'Medium List 2',
 'Medium List 2 Accent 1',
 'Medium List 2 Accent 2',
 'Medium List 2 Accent 3',
 'Medium List 2 Accent 4',
 'Medium List 2 Accent 5',
 'Medium List 2 Accent 6',
 'Medium Grid 1',
 'Medium Grid 1 Accent 1',
 'Medium Grid 1 Accent 2',
 'Medium Grid 1 Accent 3',
 'Medium Grid 1 Accent 4',
 'Medium Grid 1 Accent 5',
 'Medium Grid 1 Accent 6',
 'Medium Grid 2',
 'Medium Grid 2 Accent 1',
 'Medium Grid 2 Accent 2',
 'Medium Grid 2 Accent 3',
 'Medium Grid 2 Accent 4',
 'Medium Grid 2 Accent 5',
 'Medium Grid 2 Accent 6',
 'Medium Grid 3',
 'Medium Grid 3 Accent 1',
 'Medium Grid 3 Accent 2',
 'Medium Grid 3 Accent 3',
 'Medium Grid 3 Accent 4',
 'Medium Grid 3 Accent 5',
 'Medium Grid 3 Accent 6',
 'Dark List',
 'Dark List Accent 1',
 'Dark List Accent 2',
 'Dark List Accent 3',
 'Dark List Accent 4',
 'Dark List Accent 5',
 'Dark List Accent 6',
 'Colorful Shading',
 'Colorful Shading Accent 1',
 'Colorful Shading Accent 2',
 'Colorful Shading Accent 3',
 'Colorful Shading Accent 4',
 'Colorful Shading Accent 5',
 'Colorful Shading Accent 6',
 'Colorful List',
 'Colorful List Accent 1',
 'Colorful List Accent 2',
 'Colorful List Accent 3',
 'Colorful List Accent 4',
 'Colorful List Accent 5',
 'Colorful List Accent 6',
 'Colorful Grid',
 'Colorful Grid Accent 1',
 'Colorful Grid Accent 2',
 'Colorful Grid Accent 3',
 'Colorful Grid Accent 4',
 'Colorful Grid Accent 5',
 'Colorful Grid Accent 6']


def ensure_style(doc, style_name, base_style='Normal', add_if_not_exist=True):
    """
    Ensure that a specific style exists in the given document. If the style does not exist,
    it can optionally be added based on the `add_if_not_exist` parameter.

    Args:
        doc (Document): The document object where the style should be ensured.
        style_name (str): The name of the style to ensure.
        base_style (str, optional): The base style to use if the style needs to be created. Defaults to 'Normal'.
        add_if_not_exist (bool, optional): If True, the style will be added if it does not exist. Defaults to True.

    Returns:
        Style: The ensured or newly created style object.

    Raises:
        ValueError: If the style does not exist and `add_if_not_exist` is False.
    """
    styles = doc.styles
    try:
        style = styles[style_name]  # Try to get the style
    except KeyError:
        if add_if_not_exist:
            # If the style does not exist and add_if_not_exist is True, create it
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles[base_style]
        else:
            raise ValueError(f"Style '{style_name}' not found and not added.")
    return style


def tree_to_list(root_node):
    """
    Converts a tree structure into a list of dictionaries using pre-order traversal.

    Args:
        root_node (Node): The root node of the tree.

    Returns:
        list: A list of dictionaries where each dictionary represents a node in the tree.
              Each dictionary contains the following keys:
              - 'name': The name of the node.
              - 'parent': The parent of the node.
              - 'style': The style associated with the node.
              - 'doc_style': The document style associated with the node.
    """
    result = []
    for node in PreOrderIter(root_node):
        node_dict = {
            'name': node.name,
            'parent': node.parent,
            'style': node.style,
            'doc_style':  node.doc_style,
        }
        result.append(node_dict)
    return result

def set_style(tag_name, level):
    """
    Determines the style to be applied based on the HTML tag name and its level.

    Args:
        tag_name (str): The name of the HTML tag (e.g., 'h1', 'h2', 'ul', 'ol', 'p').
        level (int): The level or depth of the tag, used for nested lists.

    Returns:
        str or None: The corresponding style name for the given tag and level. 
                     Returns 'None' for paragraph tags ('p').
    """
    if tag_name=="h1":
        return "Heading 1"
    if tag_name=="h2":
        return "Heading 2"
    if tag_name=="h3":
        return "Heading 3"
    if tag_name=="h4":
        return "Heading 4"
    if tag_name=="h5":
        return "Heading 5"
    if tag_name=="ul" and level==1:
        return "List Bullet"
    if tag_name=="ul" and level>1:
        style = "List Bullet "+str(level)
        return style.strip()
    if tag_name=="ol" and level==1:
        return "List Number"
    if tag_name=="ol" and level>1:
        style = "List Number "+str(level)
        return style.strip()
    if tag_name=="p":
        return None 



# Function to recursively build the tree using anytree
def build_tree(bs_element, parent=None, level=0):
    """
    Recursively builds a tree structure from a BeautifulSoup element.
    Args:
        bs_element (bs4.element.Tag): The BeautifulSoup element to process.
        parent (Node, optional): The parent node in the tree. Defaults to None.
        level (int, optional): The current level in the tree hierarchy. Defaults to 0.
    Returns:
        None: The function modifies the tree in place and does not return a value.
    """
    if bs_element.name is not None:
        # ####
        # if bs_element.name =='li':
        #     direct_text = str(bs_element)
        # ####
        # else:
        direct_text = ''.join(bs_element.find_all(text=True, recursive=False)).strip()
        style_name = bs_element.name
        if style_name.strip() == 'li':
            style_name = bs_element.parent.name
        
        if bs_element.name not in ['ul', 'ol']:
            current_node = Node(name=direct_text or bs_element.name, parent=parent, level=level, style=style_name, doc_style=set_style(style_name, level))
        else:
            current_node = parent 
        for child in bs_element.children:
            if child.name and child.name in ['ul', 'ol']:
                build_tree(child, current_node, level)
            else:
                build_tree(child, current_node, level+1)
            


def get_tree_list(markdown_content):
    """
    Converts markdown content into a tree list structure.

    This function takes markdown content as input, converts it to HTML, parses the HTML to build a tree structure,
    and then converts the tree structure into a list format.

    Args:
        markdown_content (str): The markdown content to be converted.

    Returns:
        list: A list representation of the tree structure, excluding the first two elements.
    """
    html_content = markdown2.markdown(markdown_content)

    soup = BeautifulSoup(html_content, 'html.parser')

    root_node = Node('root', parent=None, level=0, style=None, doc_style=None)  # Optional: set a root node if you like

    build_tree(soup, root_node)

    tree_list = tree_to_list(root_node)

    return tree_list[2:]

def attach_paragraphs(tree_list, para):
    """
    Attaches paragraphs to a given paragraph object based on a list of nodes.

    This function iterates over a list of nodes in reverse order and inserts 
    paragraphs before the given paragraph object. If a node has a 'doc_style', 
    the paragraph is inserted with that style; otherwise, it is inserted without 
    any specific style. The iteration stops if a node with the name 'doucment' 
    (case-insensitive) is encountered.

    Args:
        tree_list (list): A list of nodes, where each node is a dictionary 
                          containing 'name' and 'doc_style' keys.
        para (Paragraph): A paragraph object before which new paragraphs will 
                          be inserted.

    Returns:
        None
    """
    for node in tree_list[::-1]:
        if node['name'].strip().lower() == str(['doucment']):
            break 
        if node['doc_style'] is not None:
            para = para.insert_paragraph_before(node['name'], style=node['doc_style'])
        else:
            para = para.insert_paragraph_before(node['name'])


def iter_block_items(parent):
    """
    Iterate over block-level items (paragraphs and tables) within a parent element.
    This function yields Paragraph or Table objects depending on the type of child elements
    found within the parent element.
    Args:
        parent (Union[doctwo, _Cell]): The parent element which can be either a doctwo or _Cell instance.
    Yields:
        Union[Paragraph, Table]: Yields Paragraph objects for paragraph elements (CT_P) and Table objects
        for table elements (CT_Tbl) found within the parent element.
    Raises:
        ValueError: If the parent element is neither a doctwo nor a _Cell instance.
    """
  
    if isinstance(parent, doctwo):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def read_template(file_name="USPI-PRL Template Word.docx"):
    """
    Reads a Word document template and ensures that specified styles are present.

    Args:
        file_name (str): The name of the Word document template to read. Defaults to "USPI-PRL Template Word.docx".

    Returns:
        Document: The Word document with ensured styles.
    """
    doc = Document(file_name)
    for style in STYLES:
        ensure_style(doc, style_name=style)
    return doc 

def extract_paragraphs(doc):
    """
    Extracts paragraphs from a given document.

    Args:
        doc: A Document object from which paragraphs are to be extracted.

    Returns:
        list: A list of paragraphs extracted from the document.
    """
    return list(iter_block_items(doc))


def extract_paragraphs_h1(doc_paragrahs, style="Heading 1"):
    """
    Extracts paragraphs from a document that match a specified style.

    Args:
        doc_paragrahs (list): A list of paragraph objects from a document.
        style (str, optional): The style to match paragraphs against. Defaults to "Heading 1".

    Returns:
        list: A list of dictionaries, each containing:
            - 'object': The paragraph object.
            - 'index': The index of the paragraph in the original list.
            - 'text': The text content of the paragraph.
    """
    paragraphs = []
    for index, paragraph in enumerate(doc_paragrahs):
        para_style = paragraph.style.name.strip().lower()
        style = style.strip().lower()
        if para_style.find(style) != -1:
            if paragraph.text.strip()!="":
                paragraphs.append({'object': paragraph, 'index': index, 'text': paragraph.text.strip()})
    return paragraphs


def update_sample_data(sample_data, template_paragraphs):
    """
    Updates the sample data with matching template paragraphs.

    This function iterates over the sample data and template paragraphs, and updates
    the sample data entries with the corresponding template paragraphs if a match is found.
    The matching is determined by the `match_string` function, which compares the text
    of the template paragraph with the section name of the sample data.

    Args:
        sample_data (list of dict): A list of dictionaries containing sample data. Each dictionary
                                    should have a 'section_name' key among other possible keys.
        template_paragraphs (list of dict): A list of dictionaries containing template paragraphs.
                                            Each dictionary should have a 'text' key among other possible keys.

    Returns:
        list of dict: The updated sample data list with the matching template paragraphs merged into it.
    """
    for index, data in enumerate(sample_data):
        for paragraph in template_paragraphs:
            if match_string(paragraph.get('text'), data['section_name']):
                sample_data[index] = {**data, **paragraph }
                template_paragraphs.remove(paragraph)
    return sample_data



def remove_number_prefix(text: str) -> str:
    # Step 1: Define a regex pattern to match numbers at the beginning, followed by optional punctuation and optional whitespace.
    pattern = r'^\d+\.\s?'
    
    # Step 2: Use re.sub() to remove the matched prefix from the text.
    result = re.sub(pattern, '', text)
    
    return result


def remove_leading_nonalpha(input_str):
    # Use regex to substitute all non-alphabet characters at the beginning with an empty string
    result = re.sub(r'^[^a-zA-Z]+', '', input_str)
    return result

def match_string(input_string: str, pattern: str) -> bool:
    input_string = remove_leading_nonalpha(input_string)
    pattern = remove_leading_nonalpha(pattern)
    # print(input_string.lower().strip())
    # print(pattern.lower().strip())
    # print('----------------')
    return input_string.lower().strip() == pattern.lower().strip()


def read_json_file(file_name):
    """
    Reads a JSON file and returns its contents as a dictionary.

    Args:
        file_name (str): The path to the JSON file to be read.

    Returns:
        dict: The contents of the JSON file.
    """
    with open(file_name) as f:
        data = json.load(f)
    return data



def save_markdown_to_html(markdown_content, filename='output.html'):
    """
    Converts Markdown content to HTML and saves it to a file.

    Args:
        markdown_content (str): The Markdown content to be converted to HTML.
        filename (str, optional): The name of the file to save the HTML content. Defaults to 'output.html'.

    Returns:
        None
    """
    html_content = markdown2.markdown(markdown_content)
    with open( filename, 'w+') as f:
        f.write(html_content)

def markdown_to_word(markdown_string, word_file):
    """
    Converts a Markdown string to a Word document.

    Args:
        markdown_string (str): The Markdown content to be converted.
        word_file (str): The file path where the Word document will be saved.

    Returns:
        None
    """
    # Convert Markdown to HTML
    html_content = markdown2.markdown(markdown_string)

    # Parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a new Word document
    doc = Document()

    # Add HTML content to the Word document
    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            p = doc.add_paragraph(element.text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListBullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListNumber')
        if element.name == 'table':
            table = doc.add_table(rows=0, cols=0)
            for row in element.find_all('tr'):
                cells = row.find_all(['td', 'th'])
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    if len(row_cells) <= i:
                        table.add_column()
                    row_cells[i].text = cell.get_text()
        else:
            doc.add_paragraph(element.text)

    # Save the Word document
    doc.save(word_file)
    return



def parse_request(sample_request):
    """
    Parses a list of request dictionaries and extracts the section name and content.

    Args:
        sample_request (list): A list of dictionaries, where each dictionary represents a request
                               and contains 'defaultSectionTitle' and 'content' keys.

    Returns:
        list: A list of dictionaries, each containing 'section_name' and 'content' keys extracted
              from the input request dictionaries.
    """
    parsed_request = []
    for request in sample_request:
        section_name = request['defaultSectionTitle']
        content = request['content']
        parsed_request.append({'section_name': section_name, 'content': content})
    return parsed_request


def move_table_after(table, paragraph):
    """
    Moves a table element in a document to immediately follow a specified paragraph.

    Args:
        table (docx.table.Table): The table to be moved.
        paragraph (docx.text.paragraph.Paragraph): The paragraph after which the table should be placed.

    Returns:
        None
    """
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)
    


def read_excel(file, doc):
    """
    Reads an Excel, JSON, or CSV file and adds its content as a table to a given Word document.
    Args:
        file (str): The path to the file to be read. The file can be in .xlsx, .json, or .csv format.
        doc (docx.Document): The Word document object where the table will be added.
    Returns:
        docx.table.Table: The table object that was added to the Word document.
    """
    import pandas as pd 
    import os 

  
    if file.endswith('.xlsx'):
        df = pd.read_excel(file)
    elif file.endswith('.json'):
        df = pd.read_json(file)
    elif file.endswith('.csv'):
        df = pd.read_csv(file)

    filename = os.path.basename(file)
    
    # doc.add_heading(filename, level=1)

    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

    for col_idx, col_name in enumerate(df.columns):
        table.cell(0, col_idx).text = col_name

    for row_idx, row in enumerate(df.values):
        for col_idx, value in enumerate(row):
            table.cell(row_idx + 1, col_idx).text = str(value)

    # table.style = 'TableGrid'

    return table


def insert_image(doc, image, paragraph):
    """
    Inserts an image into a Word document at a specified paragraph location.

    Args:
        doc (docx.Document): The Word document object where the image will be inserted.
        image (str): The file path to the image to be inserted.
        paragraph (docx.text.paragraph.Paragraph): The paragraph object before which the image will be inserted.

    Returns:
        docx.text.paragraph.Paragraph: The paragraph object where the image was inserted.
    """
    from docx.shared import Inches
    from pathlib import Path 
    file_path = "image"
    base_name = Path(file_path).stem
    p = paragraph.insert_paragraph_before('')
    r = p.add_run()
    r.add_picture(image, width=Inches(2))
    p = p.insert_paragraph_before("", 'Normal')
    # paragraph_after_p2 = paragraph.insert_paragraph_before("")  # This ensures the image goes right after p2
    # doc.add_picture(image, width=Inches(2))
    return p 


def classify_content(content):
    """
    Classifies the given content based on its type.
    Args:
        content (str): The content to classify. This can be a file path or a string.
    Returns:
        str: The classification of the content. Possible values are:
            - 'table': If the content is a file with extension .csv, .xlsx, or .json.
            - 'image': If the content is a file with extension .jpeg, .png, or .jpg.
            - 'text': If the content is not a recognized file type or is a string.
            - 'Error: <error_message>': If an exception occurs during classification.
    """
    try:
        content_path = Path(content)
        
        # Check if the content is a file path
        if content_path.is_file():
            # Get file extension
            file_extension = content_path.suffix.lower()
            
            # Check file type
            if file_extension in {'.csv', '.xlsx', '.json'}:
                return 'table'
            elif file_extension in {'.jpeg', '.png', '.jpg'}:
                return 'image'
        
        # If the content is not a recognized file type, return 'text'
        return 'text'
    except Exception as e:
        return f"Error: {str(e)}"