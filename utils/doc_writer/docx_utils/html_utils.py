from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches

import docx
import requests
import io

def parse_soup_to_elements(soup):
    """
    Processes a BeautifulSoup soup object and returns a list of Paragraph and Table objects.

    Args:
        soup: The BeautifulSoup object representing the parsed HTML content.

    Returns:
        A list of Paragraph and Table objects from python-docx.
    """
    # Create a new Word document (used as a factory for creating elements)
    document = Document()

    # List to hold Paragraph and Table objects
    elements = []

    # Process the body of the HTML
    body = soup.body or soup
    process_children(body, document, elements)

    return elements 

def process_children(parent_node, document, elements, parent_element=None):
    """
    Processes all child nodes of a parent node.

    Args:
        parent_node: The parent BeautifulSoup node.
        document: The Word Document object.
        elements: The list to store Paragraph and Table objects.
        parent_element: The parent Word element (e.g., paragraph, table cell).
    """
    for child in parent_node.contents:
        process_node(child, document, elements, parent_element)

def process_node(node, document, elements, parent_element=None):
    """
    Processes a single HTML node and converts it to Word format.

    Args:
        node: The BeautifulSoup node to process.
        document: The Word Document object.
        elements: The list to store Paragraph and Table objects.
        parent_element: The parent Word element.
    """
    try:
        if isinstance(node, NavigableString):
            text = node.strip()
            if text:
                if isinstance(parent_element, docx.text.run.Run):
                    parent_element.add_text(text)
                elif isinstance(parent_element, docx.text.paragraph.Paragraph):
                    parent_element.add_run(text)
                elif isinstance(parent_element, docx.table._Cell):
                    if parent_element.paragraphs:
                        parent_element.paragraphs[-1].add_run(text)
                    else:
                        p = parent_element.add_paragraph()
                        p.add_run(text)
                else:
                    p = document.add_paragraph()
                    p.add_run(text)
                    elements.append(p)
        elif isinstance(node, Tag):
            if node.name == 'p':
                p = document.add_paragraph()
                elements.append(p)
                process_children(node, document, elements, p)
            elif node.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                heading_level = int(node.name[1])
                style_name = f'Heading {heading_level}'
                p = document.add_paragraph(style=style_name)
                elements.append(p)
                process_children(node, document, elements, p)
            elif node.name in ['strong', 'b']:
                run = get_or_add_run(parent_element, document, elements)
                run.bold = True
                process_children(node, document, elements, run)
            elif node.name in ['em', 'i']:
                run = get_or_add_run(parent_element, document, elements)
                run.italic = True
                process_children(node, document, elements, run)
            elif node.name == 'u':
                run = get_or_add_run(parent_element, document, elements)
                run.underline = True
                process_children(node, document, elements, run)
            elif node.name == 'br':
                if isinstance(parent_element, docx.text.run.Run):
                    parent_element.add_break()
                elif isinstance(parent_element, docx.text.paragraph.Paragraph):
                    parent_element.add_run().add_break()
                elif isinstance(parent_element, docx.table._Cell):
                    if parent_element.paragraphs:
                        parent_element.paragraphs[-1].add_run().add_break()
            elif node.name in ['ul', 'ol']:
                process_list(node, document, elements, parent_element)
            elif node.name == 'table':
                process_table(node, document, elements, parent_element)
            elif node.name == 'img':
                process_image(node, document, elements, parent_element)
            elif node.name == 'a':
                process_hyperlink(node, document, elements, parent_element)
            elif node.name == 'blockquote':
                p = document.add_paragraph(style='Intense Quote')
                elements.append(p)
                process_children(node, document, elements, p)
            elif node.name == 'hr':
                # Add a horizontal line
                p = document.add_paragraph()
                run = p.add_run()
                run.add_break(docx.enum.text.WD_BREAK.LINE)
                elements.append(p)
            else:
                # Process other tags recursively
                process_children(node, document, elements, parent_element)
    except Exception as e:
        print(f"Error processing node {node}: {e}")

def get_or_add_run(parent_element, document, elements):
    """
    Gets or creates a Run object in the parent element.

    Args:
        parent_element: The parent Word element.
        document: The Word Document object.
        elements: The list to store Paragraph and Table objects.

    Returns:
        The Run object.
    """
    if isinstance(parent_element, docx.text.run.Run):
        return parent_element
    elif isinstance(parent_element, docx.text.paragraph.Paragraph):
        return parent_element.add_run()
    elif isinstance(parent_element, docx.table._Cell):
        if parent_element.paragraphs:
            return parent_element.paragraphs[-1].add_run()
        else:
            p = parent_element.add_paragraph()
            return p.add_run()
    else:
        p = document.add_paragraph()
        elements.append(p)
        return p.add_run()

def process_list(node, document, elements, parent_element=None, level=0):
    """
    Processes HTML lists and converts them to Word lists.

    Args:
        node: The BeautifulSoup node representing the list.
        document: The Word Document object.
        elements: The list to store Paragraph and Table objects.
        parent_element: The parent Word element.
        level: The nesting level of the list.
    """
    list_type = node.name
    for child in node.contents:
        if isinstance(child, Tag):
            if child.name == 'li':
                style = 'List Bullet' if list_type == 'ul' else 'List Number'
                p = document.add_paragraph(style=style)
                p.paragraph_format.left_indent = Inches(0.25 * level)
                elements.append(p)
                process_children(child, document, elements, p)
                # Check for nested lists within 'li'
                for li_child in child.contents:
                    if isinstance(li_child, Tag) and li_child.name in ['ul', 'ol']:
                        process_list(li_child, document, elements, parent_element, level+1)
            elif child.name in ['ul', 'ol']:
                process_list(child, document, elements, parent_element, level+1)

def process_table(node, document, elements, parent_element=None):
    """
    Processes HTML tables and converts them to Word tables.

    Args:
        node: The BeautifulSoup node representing the table.
        document: The Word Document object.
        elements: The list to store Paragraph and Table objects.
        parent_element: The parent Word element.
    """
    rows = node.find_all('tr')
    num_rows = len(rows)
    num_cols = max([len(r.find_all(['td', 'th'])) for r in rows]) if num_rows > 0 else 0

    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    elements.append(table)

    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            table_cell = table.cell(i, j)
            process_children(cell, document, elements, table_cell)

def process_hyperlink(node, document, elements, parent_element):
    """
    Processes HTML hyperlinks and adds them to the Word document.

    Args:
        node: The BeautifulSoup node representing the hyperlink.
        document: The Word Document object.
        elements: The list to store Paragraph and Table objects.
        parent_element: The parent Word element.
    """
    url = node.get('href', '')
    text = node.get_text(strip=True)
    if parent_element:
        if isinstance(parent_element, docx.text.paragraph.Paragraph):
            add_hyperlink(parent_element, url, text)
        elif isinstance(parent_element, docx.text.run.Run):
            # Create a new paragraph if parent is a Run
            p = document.add_paragraph()
            elements.append(p)
            add_hyperlink(p, url, text)
        else:
            p = document.add_paragraph()
            elements.append(p)
            add_hyperlink(p, url, text)
    else:
        p = document.add_paragraph()
        elements.append(p)
        add_hyperlink(p, url, text)

def add_hyperlink(paragraph, url, text):
    """
    Adds a hyperlink to a Word paragraph.

    Args:
        paragraph: The paragraph to add the hyperlink to.
        url: The URL of the hyperlink.
        text: The display text for the hyperlink.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.ns.qn('r:id'), r_id)

    new_run = docx.oxml.OxmlElement('w:r')
    rPr = docx.oxml.OxmlElement('w:rPr')

    # Style the hyperlink (blue and underlined)
    c = docx.oxml.OxmlElement('w:color')
    c.set(docx.oxml.ns.qn('w:val'), "0000FF")
    rPr.append(c)
    u = docx.oxml.OxmlElement('w:u')
    u.set(docx.oxml.ns.qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    t = docx.oxml.OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

def process_image(node, document, elements, parent_element):
    """
    Processes HTML images and adds them to the Word document.

    Args:
        node: The BeautifulSoup node representing the image.
        document: The Word Document object.
        elements: The list to store Paragraph and Table objects.
        parent_element: The parent Word element.
    """
    src = node.get('src', '')
    if src:
        try:
            # Check if the src is a URL or a local file
            if src.startswith('http'):
                response = requests.get(src)
                image_data = io.BytesIO(response.content)
            else:
                image_data = src  # Assuming src is a local file path
            if isinstance(parent_element, docx.text.run.Run):
                parent_element.add_picture(image_data)
            elif isinstance(parent_element, docx.text.paragraph.Paragraph):
                parent_element.add_run().add_picture(image_data)
            elif isinstance(parent_element, docx.table._Cell):
                parent_element.paragraphs[-1].add_run().add_picture(image_data)
            else:
                p = document.add_paragraph()
                elements.append(p)
                p.add_run().add_picture(image_data)
        except Exception as e:
            print(f"Error adding image {src}: {e}")





def insert_elements_after_paragraph(paragraph, elements):
    """
    Inserts a list of Paragraph and Table elements after a given paragraph in a document.

    Args:
        paragraph: The paragraph after which the elements will be inserted.
        elements: A list of Paragraph and Table objects to insert.
    """
    # Get the parent element (the body of the document)
    parent_element = paragraph._p.getparent()
    # Find the index of the paragraph in the parent element
    index = parent_element.index(paragraph._p)

    # Insert each element after the paragraph
    for element in elements:
        if isinstance(element, docx.text.paragraph.Paragraph):
            # Insert the paragraph's XML into the parent element
            parent_element.insert(index + 1, element._p)
            index += 1  # Move the index forward
        elif isinstance(element, docx.table.Table):
            # Insert the table's XML into the parent element
            parent_element.insert(index + 1, element._tbl)
            index += 1  # Move the index forward
        else:
            print(f"Unknown element type: {type(element)}")